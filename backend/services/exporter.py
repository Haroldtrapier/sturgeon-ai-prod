"""
Proposal DOCX exporter - generates professional proposal documents from scratch.
No template files required.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from datetime import datetime

try:
    from services.db import supabase
except ImportError:
    from backend.services.db import supabase


def export_proposal(proposal_id: str) -> str:
    """Export a proposal to a formatted DOCX file. Returns the temp file path."""
    doc = Document()

    # -- Style setup --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # -- Fetch data --
    prop_res = supabase.table("proposals") \
        .select("*, opportunities(title, agency, naics_code, response_deadline, notice_id)") \
        .eq("id", proposal_id) \
        .single() \
        .execute()
    prop = prop_res.data or {}

    sections = supabase.table("proposal_sections") \
        .select("*") \
        .eq("proposal_id", proposal_id) \
        .order("order_index") \
        .execute().data or []

    opp = prop.get("opportunities") or {}
    title = prop.get("title", "Untitled Proposal")

    # ── Cover Page ────────────────────────────────────────────
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

    if opp.get("agency"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Submitted to: {opp['agency']}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    if opp.get("notice_id"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Solicitation: {opp['notice_id']}")
        run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Date: {datetime.utcnow().strftime('%B %d, %Y')}").font.size = Pt(12)

    if opp.get("response_deadline"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Response Deadline: {opp['response_deadline']}")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)

    doc.add_page_break()

    # ── Table of Contents ─────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    for i, s in enumerate(sections, 1):
        doc.add_paragraph(
            f"{i}. {s.get('section_name', f'Section {i}')}",
            style="List Number",
        )
    doc.add_page_break()

    # ── Proposal Sections ─────────────────────────────────────
    for s in sections:
        doc.add_heading(s.get("section_name", "Untitled Section"), level=1)
        content = s.get("content", "")
        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith(("- ", "• ", "* ")):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    # ── Document Info Footer ──────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Document Information", level=2)
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Light Grid Accent 1"
    for i, (label, val) in enumerate([
        ("Proposal ID", proposal_id),
        ("Status", prop.get("status", "draft")),
        ("Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")),
        ("NAICS Code", opp.get("naics_code", "N/A")),
    ]):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = str(val)

    # ── Save ──────────────────────────────────────────────────
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name
