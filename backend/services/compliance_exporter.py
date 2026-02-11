"""
Compliance matrix DOCX exporter - generates compliance traceability matrix.
No template files required.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from datetime import datetime

try:
    from services.db import supabase
except ImportError:
    from backend.services.db import supabase


STATUS_COLORS = {
    "addressed": RGBColor(0x10, 0xB9, 0x81),   # green
    "partial": RGBColor(0xF5, 0x9E, 0x0B),      # amber
    "not_addressed": RGBColor(0xEF, 0x44, 0x44), # red
    "na": RGBColor(0x94, 0xA3, 0xB8),            # gray
}


def export_compliance(proposal_id: str) -> str:
    """Export compliance matrix to a formatted DOCX file. Returns temp file path."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # -- Fetch data --
    prop_res = supabase.table("proposals") \
        .select("title, opportunities(title, agency, notice_id)") \
        .eq("id", proposal_id) \
        .single() \
        .execute()
    prop = prop_res.data or {}
    opp = prop.get("opportunities") or {}

    reqs = supabase.table("compliance_requirements") \
        .select("*") \
        .eq("proposal_id", proposal_id) \
        .order("created_at") \
        .execute().data or []

    # ── Header ────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Compliance Traceability Matrix")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(prop.get("title", "Untitled Proposal")).font.size = Pt(12)

    if opp.get("notice_id"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Solicitation: {opp['notice_id']}").font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generated: {datetime.utcnow().strftime('%B %d, %Y')}").font.size = Pt(10)

    doc.add_paragraph("")  # spacer

    # ── Summary ───────────────────────────────────────────────
    total = len(reqs)
    addressed = sum(1 for r in reqs if r.get("status") == "addressed")
    partial = sum(1 for r in reqs if r.get("status") == "partial")
    not_addressed = sum(1 for r in reqs if r.get("status") == "not_addressed")

    doc.add_heading("Summary", level=2)
    summary_tbl = doc.add_table(rows=1, cols=4)
    summary_tbl.style = "Light Grid Accent 1"
    for i, (label, val) in enumerate([
        ("Total", str(total)),
        ("Addressed", str(addressed)),
        ("Partial", str(partial)),
        ("Not Addressed", str(not_addressed)),
    ]):
        summary_tbl.rows[0].cells[i].text = f"{label}: {val}"

    doc.add_paragraph("")

    # ── Requirements Matrix ───────────────────────────────────
    doc.add_heading("Requirements Matrix", level=2)

    if not reqs:
        doc.add_paragraph("No compliance requirements found for this proposal.")
    else:
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Light Grid Accent 1"

        headers = ["#", "Requirement", "Section Ref", "Status", "Notes"]
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for idx, r in enumerate(reqs, 1):
            row = tbl.add_row().cells
            row[0].text = str(idx)
            row[1].text = r.get("requirement", "")[:200]
            row[2].text = r.get("section_ref", "") or ""
            status = r.get("status", "not_addressed")
            row[3].text = status.replace("_", " ").title()
            row[4].text = r.get("notes", "") or ""

    # ── Save ──────────────────────────────────────────────────
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name
